from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
RESULTS = ROOT / "results"
FIGURES = ROOT / "figures"
REPORT = ROOT / "report"
OUTPUT_DOCX = REPORT / "bao_cao_dp_sgd_chi_tiet_15_trang.docx"
OUTPUT_MD = REPORT / "bao_cao_dp_sgd_chi_tiet_15_trang.md"


def fmt(value: object, digits: int = 4) -> str:
    if value is None or pd.isna(value):
        return "N/A"
    if isinstance(value, float):
        return f"{value:.{digits}f}"
    return str(value)


def add_paragraph(document: Document, text: str, style: str | None = None) -> None:
    paragraph = document.add_paragraph(style=style)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = paragraph.add_run(item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"


def add_table(document: Document, df: pd.DataFrame, columns: list[str], title: str | None = None) -> None:
    if title:
        add_paragraph(document, title)
    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for idx, col in enumerate(columns):
        cell = table.rows[0].cells[idx]
        cell.text = col
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(9)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for idx, col in enumerate(columns):
            value = row[col]
            if isinstance(value, float):
                text = fmt(value)
            else:
                text = str(value)
            cells[idx].text = text
            for paragraph in cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(8)


def add_image(document: Document, image_path: Path, width: float = 6.2) -> None:
    if image_path.exists():
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(image_path), width=Inches(width))


def load_data() -> dict[str, object]:
    grid = pd.read_csv(RESULTS / "best_variant_50epoch_noise_clip_results.csv")
    baseline = pd.read_csv(RESULTS / "baseline_results.csv")
    boosted = json.loads((RESULTS / "best_boosted_feature_baseline.json").read_text(encoding="utf-8"))
    summary = json.loads((RESULTS / "best_variant_50epoch_noise_clip_summary.json").read_text(encoding="utf-8"))
    return {
        "grid": grid,
        "baseline": baseline,
        "boosted": boosted,
        "summary": summary,
    }


def make_pivot_tables(grid: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    f1 = grid.pivot(index="noise_multiplier", columns="max_grad_norm", values="f1_score").sort_index()
    eps = grid.pivot(index="noise_multiplier", columns="max_grad_norm", values="epsilon").sort_index()
    best = (
        grid.sort_values("f1_score", ascending=False)
        .groupby("noise_multiplier", as_index=False)
        .first()
        .sort_values("noise_multiplier")
    )
    f1.columns = [f"clip={col:g}" for col in f1.columns]
    eps.columns = [f"clip={col:g}" for col in eps.columns]
    f1 = f1.reset_index().rename(columns={"noise_multiplier": "noise"})
    eps = eps.reset_index().rename(columns={"noise_multiplier": "noise"})
    return f1, eps, best


def markdown_table(df: pd.DataFrame, columns: list[str]) -> str:
    lines = ["| " + " | ".join(columns) + " |", "| " + " | ".join(["---"] * len(columns)) + " |"]
    for _, row in df.iterrows():
        lines.append("| " + " | ".join(fmt(row[col]) if isinstance(row[col], float) else str(row[col]) for col in columns) + " |")
    return "\n".join(lines)


def create_markdown(data: dict[str, object], f1: pd.DataFrame, eps: pd.DataFrame, best: pd.DataFrame) -> str:
    boosted = data["boosted"]
    grid = data["grid"]
    best_row = grid.sort_values("f1_score", ascending=False).iloc[0]
    worst_row = grid.sort_values("f1_score", ascending=True).iloc[0]
    full = grid[
        [
            "noise_multiplier",
            "max_grad_norm",
            "epsilon",
            "accuracy",
            "precision",
            "recall",
            "f1_score",
            "roc_auc",
            "pr_auc",
        ]
    ].sort_values(["noise_multiplier", "max_grad_norm"])
    return f"""# Báo cáo chi tiết DP-SGD trên UCI Adult

## Tóm tắt

Báo cáo này dùng số liệu thật từ grid cuối cùng: `mlp_64_32_d0p1`, 50 epoch, không early stopping, chạy trên NVIDIA T1200. Baseline mạnh nhất hiện tại là HistGradientBoosting với engineered features, F1 = {fmt(boosted['f1_score'])}. Kết quả DP-SGD tốt nhất trong grid 50 epoch là F1 = {fmt(best_row['f1_score'])}, noise = {fmt(best_row['noise_multiplier'], 1)}, clip = {fmt(best_row['max_grad_norm'], 1)}.

## Bảng F1 theo noise x clip

{markdown_table(f1, list(f1.columns))}

## Bảng epsilon theo noise x clip

{markdown_table(eps, list(eps.columns))}

## Best per noise

{markdown_table(best[['noise_multiplier', 'max_grad_norm', 'epsilon', 'accuracy', 'precision', 'recall', 'f1_score']], ['noise_multiplier', 'max_grad_norm', 'epsilon', 'accuracy', 'precision', 'recall', 'f1_score'])}

## Full grid

{markdown_table(full, list(full.columns))}

## Ghi chú

F1 không bắt buộc đơn điệu tuyệt đối theo từng ô, nhất là ở vùng noise rất cao và clip lớn. Điều quan trọng là xu hướng tổng thể: epsilon giảm rõ khi noise tăng, còn F1 giảm mạnh ở các cấu hình high-noise/high-clip.

Worst configuration: noise = {fmt(worst_row['noise_multiplier'], 1)}, clip = {fmt(worst_row['max_grad_norm'], 1)}, F1 = {fmt(worst_row['f1_score'])}.
"""


def build_report() -> None:
    REPORT.mkdir(exist_ok=True)
    data = load_data()
    grid: pd.DataFrame = data["grid"]
    baseline: pd.DataFrame = data["baseline"]
    boosted: dict[str, object] = data["boosted"]
    summary: dict[str, object] = data["summary"]
    f1_pivot, eps_pivot, best_by_noise = make_pivot_tables(grid)

    best_row = grid.sort_values("f1_score", ascending=False).iloc[0]
    worst_row = grid.sort_values("f1_score", ascending=True).iloc[0]
    full_grid = grid[
        [
            "noise_multiplier",
            "max_grad_norm",
            "epsilon",
            "accuracy",
            "precision",
            "recall",
            "f1_score",
            "roc_auc",
            "pr_auc",
        ]
    ].sort_values(["noise_multiplier", "max_grad_norm"])

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BÁO CÁO CHI TIẾT\nDP-SGD TRÊN UCI ADULT DATASET")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(18)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Full grid noise x clipping cho mô hình tốt nhất, 50 epoch, không early stopping")
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    document.add_paragraph()
    add_paragraph(
        document,
        "Báo cáo này được tạo từ các file kết quả trong thư mục results/. Không chỉnh sửa hoặc làm đẹp số liệu thủ công. "
        "Mục tiêu là trình bày đầy đủ cơ chế DP-SGD, cách gradient clipping và Gaussian noise được áp dụng, đồng thời phân tích grid thực nghiệm cuối cùng theo F1-score.",
    )
    document.add_page_break()

    add_heading(document, "1. Tóm tắt điều hành", 1)
    add_paragraph(
        document,
        f"Thí nghiệm cuối cùng giữ lại biến thể mô hình tốt nhất từ grid trước: {summary['model_variant']['name']} "
        f"với hidden layers {tuple(summary['model_variant']['hidden_dims'])}, dropout {summary['model_variant']['dropout']}, "
        "batch size 256, learning rate 0.05, SGD momentum 0.9 và delta = 1e-5. Mỗi cấu hình được train đủ 50 epoch, không dùng early stopping.",
    )
    add_bullets(
        document,
        [
            f"Baseline mạnh nhất không-DP: HistGradientBoosting + engineered features, F1 = {fmt(boosted['f1_score'])}, ROC-AUC = {fmt(boosted['roc_auc'])}.",
            f"DP-SGD tốt nhất trong grid 50 epoch: noise = {fmt(best_row['noise_multiplier'], 1)}, clip = {fmt(best_row['max_grad_norm'], 1)}, epsilon = {fmt(best_row['epsilon'])}, F1 = {fmt(best_row['f1_score'])}.",
            f"Cấu hình tệ nhất trong grid: noise = {fmt(worst_row['noise_multiplier'], 1)}, clip = {fmt(worst_row['max_grad_norm'], 1)}, epsilon = {fmt(worst_row['epsilon'])}, F1 = {fmt(worst_row['f1_score'])}.",
            "Kết quả cho thấy tăng noise làm epsilon giảm mạnh. Utility không giảm đơn điệu ở mọi ô, nhưng high-noise/high-clip gây suy giảm F1 rõ rệt.",
        ],
    )
    document.add_page_break()

    add_heading(document, "2. Bối cảnh bài toán và dữ liệu", 1)
    add_paragraph(
        document,
        "Bài toán sử dụng UCI Adult/Census Income để dự đoán thu nhập >50K hay <=50K. Đây là dữ liệu tabular cá nhân gồm tuổi, nghề nghiệp, học vấn, tình trạng hôn nhân, giới tính, số giờ làm việc mỗi tuần và các thuộc tính liên quan. Vì dữ liệu chứa thông tin cá nhân, bài toán phù hợp để minh họa Privacy in Development: quyền riêng tư được đưa trực tiếp vào quá trình huấn luyện thay vì chỉ xử lý sau khi mô hình đã triển khai.",
    )
    add_paragraph(
        document,
        "Pipeline tiền xử lý sử dụng split gốc của Adult: adult.data làm train và adult.test làm test. Các dòng có giá trị thiếu '?' được loại bỏ; biến phân loại được one-hot encoding; biến số được chuẩn hóa theo thống kê train. Sau xử lý, dữ liệu có 30,162 mẫu train, 15,060 mẫu test và 104 features.",
    )
    add_table(
        document,
        baseline[["model_name", "accuracy", "precision", "recall", "f1_score", "roc_auc", "pr_auc"]],
        ["model_name", "accuracy", "precision", "recall", "f1_score", "roc_auc", "pr_auc"],
        "Bảng 1. Baseline ban đầu.",
    )
    document.add_page_break()

    add_heading(document, "3. Baseline mạnh không-DP", 1)
    add_paragraph(
        document,
        "Để tránh việc DP-SGD trông tốt hơn chỉ vì baseline yếu, dự án đã thử thêm các baseline mạnh hơn: RandomForest, ExtraTrees, HistGradientBoosting, XGBoost, LightGBM, CatBoost và các feature engineering như binning tuổi/số giờ làm, capital_net, has_capital_gain, has_capital_loss, education_hours, occupation_education và marital_relationship.",
    )
    add_paragraph(
        document,
        "Ngưỡng phân loại cho baseline boosted được chọn bằng validation split nội bộ từ train, không chọn trực tiếp trên test. Điều này giúp tránh leakage trong đánh giá F1.",
    )
    boosted_df = pd.DataFrame(
        [
            {
                "model": boosted["model_name"],
                "default_accuracy": boosted["default_accuracy"],
                "default_f1": boosted["default_f1_score"],
                "tuned_accuracy": boosted["accuracy"],
                "tuned_f1": boosted["f1_score"],
                "roc_auc": boosted["roc_auc"],
                "pr_auc": boosted["pr_auc"],
            }
        ]
    )
    add_table(document, boosted_df, list(boosted_df.columns), "Bảng 2. Baseline mạnh nhất sau feature engineering.")
    add_paragraph(
        document,
        "Kết quả này là mốc tham chiếu utility không-DP chính cho phần DP-SGD. Mốc F1 = 0.7273 cao hơn đáng kể so với MLP baseline ban đầu, do đó các so sánh DP-SGD sau này không dựa vào một baseline yếu.",
    )
    document.add_page_break()

    add_heading(document, "4. Cơ chế DP-SGD", 1)
    add_paragraph(
        document,
        "DP-SGD thay đổi vòng lặp huấn luyện SGD thông thường bằng cách xử lý gradient ở cấp từng mẫu. Thay vì chỉ tính một gradient trung bình cho mini-batch rồi cập nhật mô hình, DP-SGD tính gradient riêng cho từng sample, giới hạn độ lớn từng gradient bằng clipping, thêm Gaussian noise vào tổng gradient đã clip, rồi mới cập nhật trọng số.",
    )
    add_bullets(
        document,
        [
            "Bước 1: Lấy một mini-batch từ tập train.",
            "Bước 2: Tính per-example gradient cho từng mẫu trong mini-batch.",
            "Bước 3: Clip từng gradient để norm L2 không vượt quá max_grad_norm.",
            "Bước 4: Cộng Gaussian noise vào gradient tổng/trung bình sau clipping.",
            "Bước 5: Cập nhật mô hình bằng gradient đã được bảo vệ.",
            "Bước 6: Privacy accountant tính epsilon cho delta cố định.",
        ],
    )
    add_paragraph(
        document,
        "Trong dự án, delta được cố định là 1e-5. Epsilon càng nhỏ nghĩa là privacy budget càng chặt và bảo vệ quyền riêng tư mạnh hơn. Tuy nhiên epsilon thấp thường đi kèm với việc thêm nhiều noise hơn, làm giảm utility của mô hình.",
    )
    document.add_page_break()

    add_heading(document, "5. Gradient clipping chi tiết", 1)
    add_paragraph(
        document,
        "Gradient clipping là bước giới hạn ảnh hưởng tối đa của một cá nhân lên cập nhật mô hình. Với mỗi sample i, nếu gradient g_i có norm lớn hơn ngưỡng C = max_grad_norm, gradient được co lại theo công thức: g_i_clipped = g_i * min(1, C / ||g_i||_2). Nhờ đó, không mẫu đơn lẻ nào có thể tạo ra cập nhật quá lớn.",
    )
    add_paragraph(
        document,
        "Trong grid cuối, bốn giá trị clipping được thử: 0.5, 1.0, 2.0, 3.0. Clip nhỏ như 0.5 bảo thủ hơn: nhiều gradient bị cắt mạnh, làm tín hiệu học yếu hơn nhưng ổn định hơn ở vùng noise cao. Clip lớn như 3.0 giữ lại nhiều tín hiệu gradient hơn ở noise thấp, nhưng khi noise tăng cao thì tổng update trở nên rất nhiễu, dẫn đến F1 sụp mạnh ở một số cấu hình.",
    )
    add_bullets(
        document,
        [
            "clip = 0.5: ổn định nhất ở noise cao; F1 giảm chậm hơn tại noise 8, 10, 15, 20.",
            "clip = 1.0: cân bằng ở noise thấp đến trung bình.",
            "clip = 2.0: có thể giữ utility ở noise thấp, nhưng bắt đầu mất ổn định từ noise 5 trở lên.",
            "clip = 3.0: tốt nhất ở noise 0.5 nhưng dễ collapse khi noise lớn.",
        ],
    )
    document.add_page_break()

    add_heading(document, "6. Gaussian noise chi tiết", 1)
    add_paragraph(
        document,
        "Noise multiplier điều khiển độ lớn Gaussian noise được thêm vào gradient đã clip. Trong Opacus, độ lệch chuẩn của noise tỷ lệ với noise_multiplier và max_grad_norm. Hiểu đơn giản, nếu clip norm là C và noise multiplier là sigma, noise Gaussian có scale xấp xỉ sigma * C trước khi được chuẩn hóa theo batch/update.",
    )
    add_paragraph(
        document,
        "Điều này giải thích vì sao noise và clip phải được phân tích cùng nhau. Một noise multiplier cao ở clip nhỏ có thể vẫn giữ được tín hiệu học tương đối ổn định, trong khi cùng noise multiplier đó ở clip lớn có thể làm update dao động mạnh hơn. Vì vậy báo cáo không chỉ chạy nhiều noise mà còn chạy full grid noise x clip.",
    )
    add_table(document, eps_pivot, list(eps_pivot.columns), "Bảng 3. Epsilon theo noise và clip. Epsilon không đổi theo clip khi noise, sampling rate, epoch và delta cố định.")
    document.add_page_break()

    add_heading(document, "7. Thiết kế thực nghiệm cuối cùng", 1)
    add_paragraph(
        document,
        "Sau khi chạy full model x clip x noise, biến thể tốt nhất được giữ lại để chạy lại dài hơn là mlp_64_32_d0p1. Mô hình gồm input layer 104 chiều, Linear 64, ReLU, Dropout 0.1, Linear 32, ReLU, Dropout 0.1 và output 2 logits. Mô hình được train bằng DP-SGD với SGD momentum 0.9, learning rate 0.05, batch size 256 và 50 epoch đầy đủ.",
    )
    add_bullets(
        document,
        [
            "Model variant: mlp_64_32_d0p1.",
            "Epochs: 50.",
            "Early stopping: không dùng.",
            "Noise multipliers: 0.5, 1, 2, 3, 5, 8, 10, 15, 20.",
            "Max grad norm: 0.5, 1, 2, 3.",
            "Tổng số cấu hình: 36.",
            "Metric chính: F1-score vì Adult bị lệch lớp.",
        ],
    )
    document.add_page_break()

    add_heading(document, "8. Kết quả F1 toàn bộ grid", 1)
    add_paragraph(
        document,
        "Bảng sau là bảng F1 chính của báo cáo. Mỗi dòng là một noise multiplier, mỗi cột là một max_grad_norm. Đây là số liệu thật từ file results/best_variant_50epoch_noise_clip_results.csv.",
    )
    add_table(document, f1_pivot, list(f1_pivot.columns), "Bảng 4. F1-score theo noise x clip.")
    add_image(document, FIGURES / "best_variant_50epoch_f1_heatmap.png", width=6.3)
    document.add_page_break()

    add_heading(document, "9. Best clip theo từng noise", 1)
    best_display = best_by_noise[
        ["noise_multiplier", "max_grad_norm", "epsilon", "accuracy", "precision", "recall", "f1_score"]
    ].rename(columns={"noise_multiplier": "noise", "max_grad_norm": "best_clip"})
    add_table(document, best_display, list(best_display.columns), "Bảng 5. Cấu hình tốt nhất theo từng noise.")
    add_paragraph(
        document,
        "Khi noise thấp, clip lớn hơn có thể giúp mô hình giữ được nhiều tín hiệu học hơn. Ví dụ noise 0.5 đạt F1 tốt nhất tại clip 3.0. Tuy nhiên khi noise tăng cao, clip 0.5 trở thành lựa chọn ổn định hơn. Từ noise 2.0 trở lên, best clip gần như luôn là 0.5, cho thấy clipping nhỏ giúp hạn chế update quá nhiễu.",
    )
    add_image(document, FIGURES / "best_variant_50epoch_f1_lines.png", width=6.3)
    document.add_page_break()

    add_heading(document, "10. Phân tích theo từng vùng noise", 1)
    add_paragraph(
        document,
        "Vùng noise thấp gồm noise 0.5 và 1.0. Ở vùng này epsilon còn khá lớn, privacy yếu hơn, nhưng utility cao nhất. F1 dao động quanh 0.66-0.676. Sự khác biệt giữa các clip chưa quá lớn, và clip 3.0 cho kết quả tốt nhất tại noise 0.5.",
    )
    add_paragraph(
        document,
        "Vùng noise trung bình gồm noise 2.0, 3.0 và 5.0. Tại đây epsilon giảm xuống đáng kể, nhưng F1 bắt đầu phân hóa theo clip. Clip 0.5 và 1.0 vẫn giữ được F1 tương đối tốt, trong khi clip 2.0 và 3.0 bắt đầu giảm mạnh hơn, đặc biệt ở noise 5.0.",
    )
    add_paragraph(
        document,
        "Vùng noise cao gồm noise 8.0, 10.0, 15.0 và 20.0. Đây là vùng stress test để minh họa utility collapse. Clip 0.5 vẫn giữ được F1 cao nhất trong mỗi noise, nhưng F1 cũng giảm từ 0.6301 tại noise 8 xuống 0.5558 tại noise 20. Clip 2.0 và 3.0 giảm rất mạnh, có cấu hình rơi xuống F1 khoảng 0.25-0.40.",
    )
    document.add_page_break()

    add_heading(document, "11. Vì sao F1 không đơn điệu tuyệt đối?", 1)
    add_paragraph(
        document,
        "Một điểm quan trọng khi trình bày DP-SGD là không nên khẳng định F1 luôn giảm đơn điệu ở từng ô khi noise tăng. DP-SGD là quá trình stochastic: mini-batch sampling, initialization, dropout, Gaussian noise và class imbalance đều có thể làm F1 dao động giữa các cấu hình gần nhau.",
    )
    add_paragraph(
        document,
        "Ví dụ ở clip 3.0, F1 tại noise 10 thấp hơn noise 15 và 20. Điều này không có nghĩa noise 15 hoặc 20 tốt hơn về tổng thể. Nó cho thấy vùng high-noise/high-clip đã mất ổn định; mô hình có thể dự đoán positive class quá ít hoặc quá nhiều, khiến precision/recall và F1 dao động. Vì vậy báo cáo cần nhìn xu hướng tổng thể, heatmap, và best clip theo noise, thay vì chỉ nhìn một cột đơn lẻ.",
    )
    add_bullets(
        document,
        [
            "Epsilon giảm gần như rõ ràng khi noise tăng.",
            "F1 giảm theo xu hướng tổng thể, nhưng không bắt buộc đơn điệu từng điểm.",
            "Ở noise cao, clip lớn làm training dễ collapse.",
            "Nếu cần kết luận định lượng mạnh hơn, bước tiếp theo là chạy multi-seed cho các cấu hình đại diện.",
        ],
    )
    document.add_page_break()

    add_heading(document, "12. Phân tích privacy-utility tradeoff", 1)
    add_paragraph(
        document,
        f"So với baseline không-DP mạnh nhất F1 = {fmt(boosted['f1_score'])}, cấu hình DP tốt nhất đạt F1 = {fmt(best_row['f1_score'])}, thấp hơn khoảng {fmt(boosted['f1_score'] - best_row['f1_score'])}. Đổi lại, DP-SGD cung cấp privacy accounting với epsilon cụ thể. Noise càng cao thì epsilon càng thấp, nhưng F1 giảm.",
    )
    add_paragraph(
        document,
        "Một cấu hình cân bằng cho phần thảo luận có thể là noise 3.0, clip 0.5: epsilon = 0.8348, F1 = 0.6424. Nếu ưu tiên utility hơn, noise 1.0, clip 1.0 đạt F1 = 0.6643 nhưng epsilon = 3.8217. Nếu ưu tiên privacy mạnh hơn, noise 10.0, clip 0.5 đạt epsilon = 0.2245 nhưng F1 còn 0.6307.",
    )
    add_paragraph(
        document,
        "Tùy mục tiêu dự án, nhóm có thể chọn cấu hình cân bằng hoặc cấu hình minh họa privacy mạnh. Điểm quan trọng là không chỉ báo cáo accuracy; F1, precision, recall, ROC-AUC và PR-AUC giúp mô tả utility đầy đủ hơn cho bài toán lệch lớp.",
    )
    document.add_page_break()

    add_heading(document, "13. Bảng full grid chi tiết", 1)
    add_paragraph(
        document,
        "Bảng full grid dưới đây giữ toàn bộ 36 cấu hình cuối cùng, gồm noise, clip, epsilon, accuracy, precision, recall, F1, ROC-AUC và PR-AUC. Đây là bảng dùng để kiểm tra lại toàn bộ kết luận trong báo cáo.",
    )
    add_table(document, full_grid, list(full_grid.columns), "Bảng 6. Full grid 50 epoch.")
    document.add_page_break()

    add_heading(document, "14. Giới hạn và hướng mở rộng", 1)
    add_paragraph(
        document,
        "Thí nghiệm hiện tại chỉ chạy một seed cho full 50-epoch grid. Vì DP-SGD có noise ngẫu nhiên, kết quả F1 ở vùng high-noise có thể dao động. Báo cáo đã tránh khẳng định đơn điệu tuyệt đối và tập trung vào xu hướng tổng thể. Để tăng độ chắc chắn, có thể chạy thêm multi-seed cho một số cấu hình đại diện như noise 0.5, 3.0, 10.0, 20.0 ở clip 0.5 và clip 3.0.",
    )
    add_paragraph(
        document,
        "Một giới hạn khác là mô hình DP-SGD vẫn là MLP nhỏ để tương thích với Opacus và chạy được trên laptop GPU. Baseline boosted sử dụng mô hình cây mạnh hơn nên không phải là kiến trúc tương đương tuyệt đối, nhưng nó là mốc utility không-DP tốt để tránh đánh giá DP-SGD dựa trên baseline yếu.",
    )
    add_bullets(
        document,
        [
            "Chạy multi-seed cho cấu hình đại diện.",
            "Thử target epsilon training thay vì chọn noise thủ công.",
            "Thử threshold tuning riêng cho DP model trên validation split nếu mục tiêu chính là F1.",
            "Bổ sung membership inference attack cho cấu hình best 50 epoch.",
            "So sánh thêm với các privacy accountant khác nếu môi trường hỗ trợ.",
        ],
    )
    document.add_page_break()

    add_heading(document, "15. Kết luận", 1)
    add_paragraph(
        document,
        "Báo cáo cho thấy DP-SGD có thể tích hợp vào pipeline PyTorch cho dữ liệu cá nhân dạng bảng. Khi noise multiplier tăng, epsilon giảm mạnh, thể hiện privacy budget chặt hơn. Utility đo bằng F1-score giảm theo xu hướng tổng thể, đặc biệt ở vùng high-noise/high-clip. Gradient clipping đóng vai trò rất quan trọng: clip nhỏ ổn định hơn khi noise cao, còn clip lớn có thể cho F1 tốt ở noise thấp nhưng dễ collapse khi noise tăng.",
    )
    add_paragraph(
        document,
        "Cấu hình tốt nhất trong grid 50 epoch là noise 0.5, clip 3.0 với F1 = 0.6761 nhưng epsilon cao 26.7664, nghĩa là privacy yếu. Nếu cần cân bằng privacy-utility, các cấu hình noise 3.0 hoặc 5.0 với clip 0.5 đáng thảo luận hơn vì epsilon thấp hơn đáng kể và F1 vẫn ở mức khoảng 0.64. Kết quả này phù hợp với thông điệp chính của DP-SGD: privacy mạnh hơn thường đi kèm chi phí utility, và chi phí đó phụ thuộc đồng thời vào noise, clipping, kiến trúc và số epoch.",
    )

    document.save(OUTPUT_DOCX)
    OUTPUT_MD.write_text(create_markdown(data, f1_pivot, eps_pivot, best_by_noise), encoding="utf-8")
    print("Saved:", OUTPUT_DOCX)
    print("Saved:", OUTPUT_MD)


if __name__ == "__main__":
    build_report()
